from __future__ import annotations

import json
import hashlib
from io import BytesIO
from pathlib import Path
import re
import subprocess
import sys
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageChops


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_txnmem_ccfa_docx import (  # noqa: E402
    TABLE_TITLES,
    _rasterize_svg,
    _scrub_metadata,
    build_document,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
DOC_PR_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS = {"w": W_NS, "wp": WP_NS}
FIG_MARKER = re.compile(r"\[\[FIG:([a-z_]+)\]\]")
MARKERS = (
    "[[FIG:",
    "[[TABLE:",
    "[[CLAIM:",
    "TXNMEM-AUTHOR-ANNOTATIONS",
    "BIBLIOGRAPHY",
    "TODO",
    "\\(",
    "\\)",
    "configs/txnmem_paper_references.json",
)
LOCAL_ARTIFACT_PATH = re.compile(
    r"(?i)(?:results/|file:/+|(?:^|[\\s(])[a-z]:[\\\\/]|/" + "Users/)"
)


def _test_rasterizer(svg: Path, cache: Path) -> tuple[Path, float]:
    root = ET.parse(svg).getroot()
    _, _, width, height = [float(value) for value in root.attrib["viewBox"].split()]
    target = cache / f"{svg.stem}-test.png"
    target.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (round(width * 2), round(height * 2)), "white")
    for x in range(image.width):
        image.putpixel((x, 0), (0, 0, 0))
        image.putpixel((x, image.height - 1), (0, 0, 0))
    for y in range(image.height):
        image.putpixel((0, y), (0, 0, 0))
        image.putpixel((image.width - 1, y), (0, 0, 0))
    image.save(target)
    return target, height / width


class TxnMemCcfaDocxTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.config = json.loads((ROOT / "configs/txnmem_ccfa_paper.json").read_text(encoding="utf-8"))
        cls.temporary_directory = TemporaryDirectory()
        cls.path = build_document(
            ROOT,
            Path(cls.temporary_directory.name) / Path(cls.config["target_output_path"]).name,
            raster_cache=Path(cls.temporary_directory.name) / "raster-cache",
            rasterizer=_test_rasterizer,
        )
        cls.document = Document(cls.path)
        cls.manifest = json.loads(
            (ROOT / "paper_assets/figures/manifest.json").read_text(encoding="utf-8")
        )["figures"]
        manuscript = (ROOT / cls.config["paper_source_path"]).read_text(encoding="utf-8")
        cls.figure_ids = FIG_MARKER.findall(manuscript)
        cls.table_ids = cls.config["body_table_ids"] + cls.config["appendix_table_ids"]
        cls.catalog = json.loads(
            (ROOT / "configs/txnmem_paper_references.json").read_text(encoding="utf-8")
        )["references"]
        with zipfile.ZipFile(cls.path) as archive:
            cls.parts = {name: archive.read(name) for name in archive.namelist()}
        cls.document_xml = cls.parts["word/document.xml"].decode("utf-8")

    @classmethod
    def tearDownClass(cls) -> None:
        cls.temporary_directory.cleanup()

    def test_configured_target_path_is_authoritative_but_never_test_output(self) -> None:
        self.assertFalse(Path(self.config["target_output_path"]).is_absolute())
        self.assertTrue(self.path.is_relative_to(Path(self.temporary_directory.name)))

    def test_full_release_build_is_deterministic_and_privacy_normalization_idempotent(self) -> None:
        root = Path(self.temporary_directory.name)
        first = root / "release-first.docx"
        second = root / "release-second.docx"
        cache = root / "raster-cache"
        build_document(ROOT, first, raster_cache=cache, rasterizer=_test_rasterizer)
        build_document(ROOT, second, raster_cache=cache, rasterizer=_test_rasterizer)
        first_hash = hashlib.sha256(first.read_bytes()).hexdigest()
        self.assertEqual(first_hash, hashlib.sha256(second.read_bytes()).hexdigest())

        _scrub_metadata(second)
        once = hashlib.sha256(second.read_bytes()).hexdigest()
        _scrub_metadata(second)
        self.assertEqual(once, hashlib.sha256(second.read_bytes()).hexdigest())
        self.assertEqual(first_hash, once)

    def test_generated_docx_has_paper_structure(self) -> None:
        self.assertTrue(self.path.is_file())
        self.assertTrue(zipfile.is_zipfile(self.path))
        self.assertEqual(
            len([name for name in self.parts if name.startswith("word/media/")]),
            6,
        )
        self.assertEqual(len(self.figure_ids), 6)
        self.assertEqual(set(self.figure_ids), set(self.manifest))
        self.assertEqual(len(self.table_ids), 8)
        self.assertEqual(set(self.table_ids), set(TABLE_TITLES))
        self.assertEqual([item["id"] for item in self.catalog], [f"R{number:02d}" for number in range(1, 33)])
        headings = [p.text for p in self.document.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn("1 引言", headings)
        self.assertIn("6 评估", headings)
        self.assertEqual(len(self.document.inline_shapes), len(self.figure_ids))
        self.assertEqual(len(self.document.tables), len(self.table_ids))

    def test_first_page_starts_as_anonymous_paper_not_report_cover(self) -> None:
        opening = [p.text.strip() for p in self.document.paragraphs[:6] if p.text.strip()]
        self.assertEqual(opening[0], "TxnMem：面向多 Agent 共享记忆的策略感知事务运行时")
        self.assertIn("匿名稿", opening)
        self.assertIn("摘要", opening)
        self.assertIn("TxnMem: A Policy-Aware Transactional Runtime for Shared Memory in Multi-Agent Systems", self.document_xml)
        self.assertNotIn("Prepared for", self.document_xml)

    def test_reader_facing_text_contains_no_source_handoff_or_annotation_markers(self) -> None:
        text = "\n".join(
            [p.text for p in self.document.paragraphs]
            + [cell.text for table in self.document.tables for row in table.rows for cell in row.cells]
        )
        for marker in MARKERS:
            self.assertNotIn(marker, text)
        self.assertNotIn("图：", "\n".join(p.text for p in self.document.paragraphs if "图：" in p.text))
        self.assertIsNone(LOCAL_ARTIFACT_PATH.search(text))

    def test_final_package_contains_no_rsid_session_identifiers(self) -> None:
        """Generated DOCX must not retain Word revision-session IDs in any XML part."""
        for name, payload in self.parts.items():
            if name.endswith(".xml"):
                self.assertNotIn(b"rsid", payload.lower(), name)

    def test_styles_lists_and_captions_are_word_native(self) -> None:
        normal = self.document.styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri")
        self.assertEqual(normal.font.size.pt, 11)
        self.assertEqual(self.document.styles["Heading 1"].font.size.pt, 16)
        self.assertEqual(self.document.styles["Heading 2"].font.size.pt, 13)
        self.assertEqual(self.document.styles["Heading 3"].font.size.pt, 12)
        self.assertTrue(any(p.style.name == "Caption" for p in self.document.paragraphs))
        self.assertTrue(any(p.style.name.startswith("List") for p in self.document.paragraphs))

    def test_figures_have_manifest_alt_text_and_captions(self) -> None:
        document_root = ET.fromstring(self.parts["word/document.xml"])
        drawings = document_root.findall(".//wp:docPr", NS)
        extents = document_root.findall(".//wp:extent", NS)
        self.assertEqual(len(drawings), len(self.figure_ids))
        self.assertEqual(len(extents), len(self.figure_ids))
        self.assertEqual(
            [drawing.attrib.get("descr") for drawing in drawings],
            [self.manifest[figure_id]["alt_text"] for figure_id in self.figure_ids],
        )
        for figure_id, extent in zip(self.figure_ids, extents):
            actual_ratio = int(extent.attrib["cy"]) / int(extent.attrib["cx"])
            dimensions = self.manifest[figure_id]["dimensions"]
            expected_ratio = dimensions["height"] / dimensions["width"]
            self.assertAlmostEqual(actual_ratio, expected_ratio, delta=0.002)
        captions = [p.text for p in self.document.paragraphs if p.style.name == "Caption"]
        self.assertEqual(
            [caption for caption in captions if caption.startswith("图")],
            [f"图 {number}  {self.manifest[figure_id]['caption']}"
             for number, figure_id in enumerate(self.figure_ids, start=1)],
        )

    def test_figure_rasters_fill_their_declared_canvas(self) -> None:
        """Catch 1× SVG content inside a 2× screenshot canvas before it becomes unreadable."""
        for name, payload in self.parts.items():
            if not name.startswith("word/media/") or not name.endswith(".png"):
                continue
            with Image.open(BytesIO(payload)).convert("RGB") as image:
                background = Image.new("RGB", image.size, "white")
                bbox = ImageChops.difference(image, background).getbbox()
            self.assertIsNotNone(bbox, name)
            content_width = bbox[2] - bbox[0]
            content_height = bbox[3] - bbox[1]
            self.assertGreaterEqual(content_width / image.width, 0.85, name)
            self.assertGreaterEqual(content_height / image.height, 0.85, name)

    def test_figure_paragraphs_keep_their_captions_together(self) -> None:
        document_root = ET.fromstring(self.parts["word/document.xml"])
        paragraphs = document_root.findall(".//w:p", NS)
        figure_paragraphs = [
            (index, paragraph) for index, paragraph in enumerate(paragraphs)
            if paragraph.find(".//w:drawing", NS) is not None
        ]
        self.assertEqual(len(figure_paragraphs), len(self.figure_ids))
        for index, paragraph in figure_paragraphs:
            keep_next = paragraph.find("./w:pPr/w:keepNext", NS)
            self.assertIsNotNone(keep_next)
            caption_keep_next = paragraphs[index + 1].find("./w:pPr/w:keepNext", NS)
            self.assertIsNotNone(caption_keep_next)
            self.assertEqual(caption_keep_next.attrib[f"{{{W_NS}}}val"], "0")

    def test_tables_have_fixed_geometry_repeating_headers_and_titles(self) -> None:
        document_root = ET.fromstring(self.parts["word/document.xml"])
        tables = document_root.findall(".//w:tbl", NS)
        self.assertEqual(len(tables), len(self.table_ids))
        repeated_headers = 0
        for table in tables:
            tbl_width = table.find("./w:tblPr/w:tblW", NS)
            self.assertIsNotNone(tbl_width)
            self.assertEqual(tbl_width.attrib[f"{{{W_NS}}}w"], "9360")
            widths = table.findall("./w:tblGrid/w:gridCol", NS)
            self.assertTrue(widths)
            self.assertTrue(all(int(width.attrib[f"{{{W_NS}}}w"]) > 0 for width in widths))
            self.assertEqual(sum(int(width.attrib[f"{{{W_NS}}}w"]) for width in widths), 9360)
            for cell in table.findall(".//w:tc", NS):
                tc_width = cell.find("./w:tcPr/w:tcW", NS)
                self.assertIsNotNone(tc_width)
                self.assertIn(int(tc_width.attrib[f"{{{W_NS}}}w"]),
                              [int(width.attrib[f"{{{W_NS}}}w"]) for width in widths])
            first_row = table.find("./w:tr", NS)
            self.assertIsNotNone(first_row)
            if first_row.find("./w:trPr/w:tblHeader", NS) is not None:
                repeated_headers += 1
            self.assertTrue(all(
                row.find("./w:trPr/w:cantSplit", NS) is not None
                for row in table.findall("./w:tr", NS)
            ))
        self.assertEqual(repeated_headers, len(tables))
        captions = [p.text for p in self.document.paragraphs if p.style.name == "Caption"]
        self.assertEqual(
            [caption for caption in captions if caption.startswith("表")],
            [f"表 {number}  {TABLE_TITLES[table_id]}"
             for number, table_id in enumerate(self.table_ids, start=1)],
        )
        self.assertEqual(len(captions), len(self.figure_ids) + len(self.table_ids))

    def test_appendix_schema_table_uses_compact_readable_type(self) -> None:
        """Keep the final appendix table and its explanatory paragraph on one page flow."""
        for number, appendix_table in ((7, self.document.tables[-2]), (8, self.document.tables[-1])):
            caption = next(
                paragraph for paragraph in self.document.paragraphs
                if paragraph.style.name == "Caption" and paragraph.text.startswith(f"表 {number}")
            )
            header_sizes = [
                run.font.size.pt
                for cell in appendix_table.rows[0].cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
            ]
            body_sizes = [
                run.font.size.pt
                for row in appendix_table.rows[1:]
                for cell in row.cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
            ]
            self.assertTrue(header_sizes and body_sizes)
            self.assertGreaterEqual(min(run.font.size.pt for run in caption.runs), 9.0)
            self.assertLessEqual(max(run.font.size.pt for run in caption.runs), 9.0)
            self.assertLessEqual(caption.paragraph_format.space_before.pt, 2.0)
            self.assertLessEqual(caption.paragraph_format.space_after.pt, 2.0)
            self.assertAlmostEqual(caption.paragraph_format.line_spacing, 220 / 240)
            self.assertGreaterEqual(min(header_sizes), 9.0)
            self.assertLessEqual(max(header_sizes), 9.0)
            self.assertGreaterEqual(min(body_sizes), 8.75)
            self.assertLessEqual(max(body_sizes), 9.0)
            for cell in [cell for row in appendix_table.rows for cell in row.cells]:
                cell_margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
                self.assertIsNotNone(cell_margins)
                self.assertEqual(cell_margins.find(qn("w:top")).get(qn("w:w")), "0")
                self.assertEqual(cell_margins.find(qn("w:bottom")).get(qn("w:w")), "0")
                self.assertEqual(cell_margins.find(qn("w:start")).get(qn("w:w")), "120")
                self.assertEqual(cell_margins.find(qn("w:end")).get(qn("w:w")), "120")
            for cell in appendix_table.rows[0].cells:
                self.assertEqual(cell.paragraphs[0].paragraph_format.line_spacing, 1.0)
            for row in appendix_table.rows[1:]:
                for cell in row.cells:
                    self.assertEqual(cell.paragraphs[0].paragraph_format.line_spacing, 1.0)
        claim_ledger_grid = self.document.tables[-2]._tbl.tblGrid.gridCol_lst
        self.assertEqual(
            [int(column.get(qn("w:w"))) for column in claim_ledger_grid],
            [1440, 1440, 6480],
        )
        appendix_note = self.document.paragraphs[-1]
        self.assertEqual(appendix_note.style.name, "TxnMem Appendix Note")
        appendix_note_style = appendix_note.style
        self.assertEqual(appendix_note_style.font.size.pt, 9.5)
        self.assertEqual(appendix_note_style.paragraph_format.space_after.pt, 0)
        self.assertAlmostEqual(appendix_note_style.paragraph_format.line_spacing, 1.0)

    def test_references_are_complete_and_stably_ordered(self) -> None:
        text = "\n".join(p.text for p in self.document.paragraphs)
        self.assertIn("参考文献", text)
        expected = [
            f"[{item['id']}] " + "; ".join(item["authors"]) + ". "
            f"{item['title']}. {item['venue']}, {item['year']}. {item['url']}"
            for item in sorted(self.catalog, key=lambda item: item["id"])
        ]
        actual = [p.text for p in self.document.paragraphs if p.text.startswith("[R")]
        self.assertEqual(actual, expected)

    def test_page_field_and_anonymous_metadata_are_present(self) -> None:
        footer_xml = b"".join(content for name, content in self.parts.items() if name.startswith("word/footer"))
        self.assertIn(b"PAGE", footer_xml)
        core = self.parts["docProps/core.xml"].decode("utf-8")
        self.assertNotIn("xiaoyan", core.lower())
        self.assertNotIn("agent-db", core.lower())
        core_root = ET.fromstring(core)
        creator = core_root.find("{http://purl.org/dc/elements/1.1/}creator")
        modified_by = core_root.find(
            "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy"
        )
        self.assertIsNone(creator)
        self.assertIsNone(modified_by)
        app = ET.fromstring(self.parts["docProps/app.xml"])
        self.assertIsNone(app.find("{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Company"))


class DocxRasterPortabilityTests(unittest.TestCase):
    def test_render_wrapper_requires_injected_version_independent_runtime_paths(self) -> None:
        script = (ROOT / "scripts/render_docx_with_bundled_libs.sh").read_text(
            encoding="utf-8"
        )
        self.assertIn('TXNMEM_CODEX_DEPS="${TXNMEM_CODEX_DEPS:-}"', script)
        self.assertIn('TXNMEM_RENDERER="${TXNMEM_RENDERER:-}"', script)
        self.assertNotIn("/" + "Users/", script)
        self.assertNotRegex(script, r"documents/\d+\.\d+\.\d+")

    def test_browser_and_raster_cache_are_injected_without_repo_writes(self) -> None:
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            svg = temp / "figure.svg"
            svg.write_text(
                '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 10 5"></svg>',
                encoding="utf-8",
            )
            browser = temp / "browser"
            browser.write_text("fixture", encoding="utf-8")
            cache = temp / "cache"
            observed: dict[str, object] = {}

            class FakeProcess:
                pid = 1234

                def poll(self):
                    return None

                def wait(self, timeout=None):
                    return 0

            def fake_popen(command, **kwargs):
                observed["command"] = command
                observed["kwargs"] = kwargs
                screenshot = next(
                    Path(item.split("=", 1)[1])
                    for item in command
                    if item.startswith("--screenshot=")
                )
                screenshot.parent.mkdir(parents=True, exist_ok=True)
                Image.new("RGB", (20, 10), "white").save(screenshot)
                return FakeProcess()

            with patch("build_txnmem_ccfa_docx.subprocess.Popen", side_effect=fake_popen), patch(
                "build_txnmem_ccfa_docx.os.killpg"
            ):
                raster, ratio = _rasterize_svg(
                    svg, cache, browser_executable=browser
                )

        self.assertEqual(raster.parent, cache)
        self.assertEqual(ratio, 0.5)
        self.assertIn("--headless=new", observed["command"])
        self.assertEqual(observed["kwargs"]["stdout"], subprocess.DEVNULL)
        self.assertNotEqual(observed["kwargs"]["stderr"], subprocess.PIPE)
        self.assertTrue(observed["kwargs"]["start_new_session"])

    def test_browser_timeout_accepts_only_a_complete_exact_raster(self) -> None:
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            svg = temp / "figure.svg"
            svg.write_text(
                '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 10 5"></svg>',
                encoding="utf-8",
            )
            browser = temp / "browser"
            browser.write_text("fixture", encoding="utf-8")

            class ReadyProcess:
                pid = 1234

                def poll(self):
                    return None

                def wait(self, timeout=None):
                    return 0

            def ready_after_screenshot(command, **kwargs):
                screenshot = next(
                    Path(item.split("=", 1)[1])
                    for item in command
                    if item.startswith("--screenshot=")
                )
                screenshot.parent.mkdir(parents=True, exist_ok=True)
                Image.new("RGB", (20, 10), "white").save(screenshot)
                return ReadyProcess()

            with patch("build_txnmem_ccfa_docx.subprocess.Popen", side_effect=ready_after_screenshot), patch(
                "build_txnmem_ccfa_docx.os.killpg"
            ):
                raster, ratio = _rasterize_svg(
                    svg, temp / "complete", browser_executable=browser
                )
            self.assertTrue(raster.is_file())
            self.assertEqual(ratio, 0.5)

            class FailedProcess:
                pid = 5678

                def poll(self):
                    return 1

                def wait(self, timeout=None):
                    return 1

            with patch("build_txnmem_ccfa_docx.subprocess.Popen", return_value=FailedProcess()):
                with self.assertRaises(RuntimeError):
                    _rasterize_svg(
                        svg, temp / "missing", browser_executable=browser
                    )

    def test_rasterizer_stops_an_isolated_browser_after_exact_raster_is_ready(self) -> None:
        """A headless browser that leaves children alive must not cost 10 s per figure."""
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            svg = temp / "figure.svg"
            svg.write_text(
                '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 10 5"></svg>',
                encoding="utf-8",
            )
            browser = temp / "browser"
            browser.write_text("fixture", encoding="utf-8")
            observed: dict[str, object] = {}

            class FakeProcess:
                pid = 4321

                def poll(self):
                    return None

                def wait(self, timeout=None):
                    observed["wait_timeout"] = timeout
                    return 0

            def fake_popen(command, **kwargs):
                observed["command"] = command
                observed["kwargs"] = kwargs
                screenshot = next(
                    Path(item.split("=", 1)[1])
                    for item in command
                    if item.startswith("--screenshot=")
                )
                screenshot.parent.mkdir(parents=True, exist_ok=True)
                Image.new("RGB", (20, 10), "white").save(screenshot)
                return FakeProcess()

            with patch("build_txnmem_ccfa_docx.subprocess.Popen", side_effect=fake_popen), patch(
                "build_txnmem_ccfa_docx.os.killpg"
            ) as killpg:
                raster, ratio = _rasterize_svg(
                    svg, temp / "cache", browser_executable=browser
                )
                self.assertTrue(raster.is_file())
                self.assertEqual(ratio, 0.5)
                self.assertTrue(observed["kwargs"]["start_new_session"])
                killpg.assert_called_once()


if __name__ == "__main__":
    unittest.main()
