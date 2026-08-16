from __future__ import annotations

import json
import hashlib
import importlib.util
from io import BytesIO
import subprocess
import sys
from tempfile import TemporaryDirectory
import unittest
from pathlib import Path
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/txnmem_paper_work_and_experiment_report_zh.md"
BUILDER = ROOT / "scripts/build_txnmem_paper_work_report_docx.py"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS = {"w": W_NS, "wp": WP_NS}


def _test_rasterizer(svg: Path, cache: Path) -> tuple[Path, float]:
    root = ET.parse(svg).getroot()
    _, _, width, height = [float(value) for value in root.attrib["viewBox"].split()]
    target = cache / f"{svg.stem}-test.png"
    target.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (round(width * 2), round(height * 2)), "white")
    for x in range(image.width):
        image.putpixel((x, 0), (0, 0, 0))
        image.putpixel((x, image.height - 1), (0, 0, 0))
    for y in range(image.height):
        image.putpixel((0, y), (0, 0, 0))
        image.putpixel((image.width - 1, y), (0, 0, 0))
    image.save(target)
    return target, height / width


def _load_builder():
    if not BUILDER.exists():
        return None
    spec = importlib.util.spec_from_file_location("txnmem_paper_work_report_builder", BUILDER)
    if spec is None or spec.loader is None:
        raise AssertionError("unable to load report builder")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class PaperWorkReportSourceTests(unittest.TestCase):
    def _source_text(self) -> str:
        self.assertTrue(SOURCE.exists(), "the versioned report source must exist")
        return SOURCE.read_text(encoding="utf-8")

    def test_report_source_covers_the_requested_narrative(self) -> None:
        text = self._source_text()
        for heading in (
            "# 执行摘要",
            "# 1. 论文研究问题与定位",
            "# 2. 主要创新点",
            "# 3. 系统与工程实现",
            "# 4. 数据构建、来源与规模",
            "# 5. 实验设计与实验矩阵",
            "# 6. 各项实验、目的与结果",
            "# 7. 创新点如何由实验闭环验证",
            "# 8. 证据治理与可复现性",
            "# 9. 结论边界、局限与投稿状态",
        ):
            self.assertIn(heading, text)

        for required_phrase in (
            "8 个 workload family",
            "50 个 seed",
            "400 个 instance",
            "2,000 条 variant row",
            "50 个 task episode",
            "110 个 native event",
            "50 个唯一 task",
            "497 个 native event",
            "20 个配对 task",
            "每次 1,986 个问题",
            "5 个场景 × 30 次重复 = 150 次观测",
            "1,632/1,632",
            "3,251,506",
            "15 条 active claim",
            "163/163",
        ):
            self.assertIn(required_phrase, text)

    def test_report_uses_active_evidence_and_preserves_claim_boundaries(self) -> None:
        text = self._source_text()
        ledger = json.loads((ROOT / "configs/paper_claims.json").read_text(encoding="utf-8"))
        active_paths = {
            claim["artifact_path"]
            for claim in ledger["claims"]
            if claim.get("status") == "active"
        }
        for path in (
            "results/paper_evidence/controlled_suite.json",
            "results/submission_evidence/tau_bench_50/aggregate.json",
            "results/submission_evidence/toxiproxy_state_verified_30/aggregate.json",
            "results/cross_host_model_load_formal_v8_aggregate/results/model_load_repetition_summary.json",
        ):
            self.assertIn(path, active_paths)
            self.assertIn(path, text)

        for boundary in (
            "reward 不是 memory accuracy",
            "不是原生 Agent memory ground truth",
            "不支持一般分布式事务、跨主机容错、可用性、线性一致性或生产延迟",
            "不是多主机 Agent workers",
            "货币成本未计算",
        ):
            self.assertIn(boundary, text)

        self.assertNotIn(
            "results/real_backend_performance_reps30_v2/results/backend_performance.json",
            text,
        )


class PaperWorkReportBuilderContractTests(unittest.TestCase):
    def test_builder_file_exists_before_release_build(self) -> None:
        self.assertTrue(BUILDER.exists(), "the deterministic DOCX builder must exist")

    def test_builder_cli_can_be_invoked_directly(self) -> None:
        completed = subprocess.run(
            [sys.executable, str(BUILDER), "--help"],
            cwd=ROOT,
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(completed.returncode, 0, completed.stderr)
        self.assertIn("--output", completed.stdout)

    def test_builder_is_deterministic_and_preserves_report_structure(self) -> None:
        self.assertTrue(BUILDER.exists(), "the deterministic DOCX builder must exist")
        builder = _load_builder()
        self.assertIsNotNone(builder)
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            first = builder.build_document(
                ROOT,
                temp / "first.docx",
                raster_cache=temp / "raster-cache",
                rasterizer=_test_rasterizer,
            )
            second = builder.build_document(
                ROOT,
                temp / "second.docx",
                raster_cache=temp / "raster-cache",
                rasterizer=_test_rasterizer,
            )
            self.assertEqual(
                hashlib.sha256(first.read_bytes()).hexdigest(),
                hashlib.sha256(second.read_bytes()).hexdigest(),
            )
            document = Document(second)
            self.assertEqual(
                next(p.text for p in document.paragraphs if p.text.strip()),
                "TxnMem 论文工作与实验总报告",
            )
            headings = [
                p.text for p in document.paragraphs
                if p.style.name.startswith("Heading")
            ]
            self.assertIn("2. 主要创新点", headings)
            self.assertIn("6. 各项实验、目的与结果", headings)
            self.assertIn("9. 结论边界、局限与投稿状态", headings)
            self.assertEqual(len(document.inline_shapes), 5)
            self.assertEqual(len(document.tables), 6)

    def test_generated_package_has_native_styles_geometry_alt_text_and_privacy(self) -> None:
        self.assertTrue(BUILDER.exists(), "the deterministic DOCX builder must exist")
        builder = _load_builder()
        self.assertIsNotNone(builder)
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            path = builder.build_document(
                ROOT,
                temp / "report.docx",
                raster_cache=temp / "raster-cache",
                rasterizer=_test_rasterizer,
            )
            document = Document(path)
            self.assertEqual(document.styles["Normal"].font.name, "Arial Unicode MS")
            self.assertEqual(document.styles["Normal"].font.size.pt, 11)
            self.assertEqual(document.styles["Heading 1"].font.size.pt, 16)
            self.assertEqual(document.styles["Heading 2"].font.size.pt, 13)
            self.assertEqual(document.styles["Heading 3"].font.size.pt, 12)
            normal_fonts = document.styles["Normal"]._element.get_or_add_rPr().get_or_add_rFonts()
            self.assertEqual(normal_fonts.get(qn("w:eastAsia")), "Arial Unicode MS")
            artifact_paragraphs = [
                paragraph for paragraph in document.paragraphs
                if paragraph.text.startswith("正式 artifact 为")
            ]
            self.assertEqual(len(artifact_paragraphs), 8)
            self.assertTrue(all(p.style.name == "TxnMem Artifact" for p in artifact_paragraphs))
            self.assertTrue(
                all(p.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT for p in artifact_paragraphs)
            )

            numbered = [p for p in document.paragraphs if p.style.name == "List Number"]
            self.assertEqual(len(numbered), 9)
            number_ids = [
                int(p._p.pPr.numPr.numId.val)
                for p in numbered
            ]
            self.assertEqual(len(set(number_ids[:3])), 1)
            self.assertEqual(len(set(number_ids[3:])), 1)
            self.assertNotEqual(number_ids[0], number_ids[3])
            with zipfile.ZipFile(path) as archive:
                parts = {name: archive.read(name) for name in archive.namelist()}
            document_root = ET.fromstring(parts["word/document.xml"])
            drawings = document_root.findall(".//wp:docPr", NS)
            self.assertEqual(len(drawings), 5)
            self.assertTrue(all(item.attrib.get("descr") for item in drawings))
            for table in document_root.findall(".//w:tbl", NS):
                grid = [int(col.attrib[qn("w:w")]) for col in table.findall("./w:tblGrid/w:gridCol", NS)]
                self.assertEqual(sum(grid), 9360)
                self.assertTrue(all(width >= 720 for width in grid))
                indent = table.find("./w:tblPr/w:tblInd", NS)
                self.assertIsNotNone(indent)
                self.assertEqual(indent.attrib[qn("w:w")], "120")
            for name, payload in parts.items():
                if name.endswith(".xml"):
                    self.assertNotIn(b"rsid", payload.lower(), name)
            core = parts["docProps/core.xml"].decode("utf-8")
            self.assertNotIn("creator", core.lower())
            self.assertNotIn("lastmodifiedby", core.lower())
            text = "\n".join(
                [p.text for p in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            for marker in ("[[FIG:", "TODO", "/" + "Users/", "TXNMEM-AUTHOR-ANNOTATIONS"):
                self.assertNotIn(marker, text)

    def test_report_has_no_unresolved_artifact_tokens_or_embedded_images_with_broken_bytes(self) -> None:
        self.assertTrue(BUILDER.exists(), "the deterministic DOCX builder must exist")
        builder = _load_builder()
        self.assertIsNotNone(builder)
        with TemporaryDirectory() as tmp:
            temp = Path(tmp)
            path = builder.build_document(
                ROOT,
                temp / "report.docx",
                raster_cache=temp / "raster-cache",
                rasterizer=_test_rasterizer,
            )
            with zipfile.ZipFile(path) as archive:
                images = [
                    archive.read(name)
                    for name in archive.namelist()
                    if name.startswith("word/media/")
                ]
            self.assertEqual(len(images), 5)
            for payload in images:
                with Image.open(BytesIO(payload)) as image:
                    image.verify()


if __name__ == "__main__":
    unittest.main()
