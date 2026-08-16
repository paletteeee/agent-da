#!/usr/bin/env python3
"""Build the Chinese TxnMem paper-work and experiment report as deterministic DOCX."""

from __future__ import annotations

import argparse
import re
import sys
import tempfile
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.build_txnmem_ccfa_docx import (
    _add_page_field,
    _plain,
    _rasterize_svg,
    _scrub_metadata,
    _set_cell_shading,
    _set_repeat_header,
    _set_run_font as _base_set_run_font,
    _set_style_font as _base_set_style_font,
    _set_table_geometry,
    _spacing,
)


TITLE = "TxnMem 论文工作与实验总报告"
SUBTITLE = "面向多 Agent 共享记忆的策略感知事务运行时"
SOURCE = Path("docs/txnmem_paper_work_and_experiment_report_zh.md")
CONTENT_WIDTH_DXA = 9360
FIGURE_WIDTH_IN = 5.5
REPORT_FONT = "Arial Unicode MS"
FIGURE_MARKER = re.compile(r"^\[\[FIG:([a-z_]+)\]\]$")
HEADING = re.compile(r"^(#{1,3})\s+(.+?)\s*$")
NUMBERED_ITEM = re.compile(r"^\d+\.\s+(.+)$")
TABLE_SEPARATOR = re.compile(r"^:?-{3,}:?$")

FIGURES = {
    "evidence_layers": (
        "paper_assets/figures/evidence_layers.svg",
        "TxnMem 分层证据结构，展示受控语义、真实模型、公开运行时、真实服务和跨主机模型服务。",
    ),
    "architecture": (
        "paper_assets/figures/architecture.svg",
        "TxnMem 系统架构，包括事务管理、策略引擎、来源管理、不变量检查和独立参考语义。",
    ),
    "commit_protocol": (
        "paper_assets/figures/commit_protocol.svg",
        "TxnMem 提交协议，从写集与依赖验证到最新策略重验证、提交或中止。",
    ),
    "provenance_repair": (
        "paper_assets/figures/provenance_repair.svg",
        "来源驱动修复流程，源对象失效后沿依赖闭包使下游对象退出默认可见集合。",
    ),
    "controlled_results": (
        "paper_assets/figures/controlled_results.svg",
        "完整 TxnMem 与四个对照在四百个受控实例上的目标违规数量比较。",
    ),
}


def _set_run_font(
    run,
    *,
    name: str = REPORT_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
) -> None:
    """Apply the report font pair with a LibreOffice-visible CJK fallback."""
    _base_set_run_font(run, name=name, size=size, color=color, bold=bold)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{script}"), REPORT_FONT)


def _set_style_font(
    style,
    *,
    size: float,
    color: str = "000000",
    bold: bool = False,
) -> None:
    _base_set_style_font(style, size=size, color=color, bold=bold)
    style.font.name = REPORT_FONT
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for script in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{script}"), REPORT_FONT)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(header.add_run("TxnMem · 论文工作与实验总报告"), size=9, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("第 "), size=9, color="667085")
    _add_page_field(footer)
    _set_run_font(footer.add_run(" 页"), size=9, color="667085")


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(normal.paragraph_format, 0, 6, 264)

    for name, values in (
        ("Heading 1", (16, "2E74B5", 16, 8)),
        ("Heading 2", (13, "2E74B5", 12, 6)),
        ("Heading 3", (12, "1F4D78", 8, 4)),
    ):
        size, color, before, after = values
        style = document.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        _spacing(style.paragraph_format, before, after, 240)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    _set_style_font(caption, size=9.5, color="475467")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(caption.paragraph_format, 3, 7, 240)
    caption.paragraph_format.keep_with_next = False

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        _set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.2)
        _spacing(style.paragraph_format, 0, 4, 264)

    meta = document.styles.add_style("TxnMem Report Meta", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(meta, size=11, color="475467")
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(meta.paragraph_format, 2, 6, 264)

    note = document.styles.add_style("TxnMem Report Note", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(note, size=10, color="344054")
    note.paragraph_format.left_indent = Inches(0.25)
    note.paragraph_format.right_indent = Inches(0.25)
    _spacing(note.paragraph_format, 4, 6, 252)

    artifact = document.styles.add_style("TxnMem Artifact", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(artifact, size=8.5, color="344054")
    artifact.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    artifact.paragraph_format.keep_together = True
    _spacing(artifact.paragraph_format, 2, 5, 240)


def _set_paragraph_bottom_border(paragraph, color: str = "2E74B5", size: str = "18") -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _append_cover(document: Document, metadata: dict[str, str]) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True
    _set_paragraph_bottom_border(title)
    _set_run_font(title.add_run(TITLE), size=26, color="1F4D78", bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(54)
    _set_run_font(
        subtitle.add_run(metadata.get("副标题", SUBTITLE)),
        size=15,
        color="344054",
        bold=True,
    )

    for label in ("报告日期", "报告范围"):
        paragraph = document.add_paragraph(style="TxnMem Report Meta")
        _set_run_font(paragraph.add_run(f"{label}："), size=11, color="667085", bold=True)
        _set_run_font(paragraph.add_run(metadata.get(label, "")), size=11, color="475467")

    scope_note = document.add_paragraph(style="TxnMem Report Note")
    scope_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_shading_for_paragraph(scope_note, "F2F4F7")
    scope_note.add_run(
        "本报告统一梳理论文问题、创新、系统实现、数据来源与规模、实验目的、主要结果、证据边界及投稿状态。"
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _set_cell_shading_for_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _parse_cover(lines: list[str]) -> tuple[dict[str, str], int]:
    metadata: dict[str, str] = {}
    index = 1
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped == "# 执行摘要":
            return metadata, index
        if "：" in stripped:
            key, value = stripped.split("：", 1)
            if key in {"副标题", "报告日期", "报告范围"}:
                metadata[key] = value.strip()
        index += 1
    raise ValueError("report source must contain '# 执行摘要'")


def _split_table_row(line: str) -> list[str]:
    return [_plain(value.strip()) for value in line.strip().strip("|").split("|")]


def _parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    block: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        block.append(_split_table_row(lines[index]))
        index += 1
    if len(block) < 2 or not all(TABLE_SEPARATOR.fullmatch(value) for value in block[1]):
        raise ValueError(f"invalid Markdown table beginning at line {start + 1}")
    return block[0], block[2:], index


def _column_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    weights = []
    for column, header in enumerate(headers):
        observed = [len(header)] + [len(row[column]) for row in rows if column < len(row)]
        weights.append(max(4, min(max(observed), 80)))
    total = sum(weights)
    minimum = 720
    distributable = CONTENT_WIDTH_DXA - minimum * len(weights)
    widths = [minimum + int(distributable * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _new_list_number_instance(document: Document) -> int:
    """Create a numbering instance that restarts at one for one Markdown block."""
    style = document.styles["List Number"]
    style_num_id = int(
        style._element.find(qn("w:pPr"))
        .find(qn("w:numPr"))
        .find(qn("w:numId"))
        .get(qn("w:val"))
    )
    numbering = document.part.numbering_part.element
    base = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == style_num_id
    )
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max(int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))) + 1

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    number.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    number.append(override)
    numbering.append(number)
    return new_id


def _apply_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    num_pr = properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        properties.append(num_pr)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), str(num_id))
    num_pr.extend((level, identifier))


def _format_table_cell(cell, value: str, *, header: bool) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(value)
    _set_run_font(run, size=8.5 if header else 8.25, bold=header)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        _set_cell_shading(cell, "F2F4F7")


def _append_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        _format_table_cell(cell, value, header=True)
    _set_repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            _format_table_cell(cell, values[index] if index < len(values) else "", header=False)
    _set_table_geometry(table, _column_widths(headers, rows), compact=True)


def _append_figure(
    document: Document,
    root: Path,
    figure_id: str,
    raster_cache: Path,
    rasterizer: Callable[[Path, Path], tuple[Path, float]],
) -> None:
    if figure_id not in FIGURES:
        raise ValueError(f"unknown figure marker: {figure_id}")
    source_name, alt_text = FIGURES[figure_id]
    source = root / source_name
    raster, ratio = rasterizer(source, raster_cache)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    # This python-docx build accepts a string path or stream, but treats Path as
    # a stream and attempts ``seek``. Normalize explicitly for both test and
    # production rasterizers.
    shape = run.add_picture(
        str(raster),
        width=Inches(FIGURE_WIDTH_IN),
        height=Inches(FIGURE_WIDTH_IN * ratio),
    )
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", f"TxnMem {figure_id}")


def _append_paragraph(
    document: Document,
    text: str,
    *,
    style: str | None = None,
    num_id: int | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    if num_id is not None:
        _apply_numbering(paragraph, num_id)
    run = paragraph.add_run(_plain(text))
    _set_run_font(run, size=11 if style is None else None)


def _append_markdown(
    document: Document,
    lines: list[str],
    start: int,
    root: Path,
    raster_cache: Path,
    rasterizer: Callable[[Path, Path], tuple[Path, float]],
) -> None:
    index = start
    active_number_id: int | None = None
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            active_number_id = None
            index += 1
            continue

        numbered = NUMBERED_ITEM.fullmatch(stripped)
        if numbered:
            if active_number_id is None:
                active_number_id = _new_list_number_instance(document)
            _append_paragraph(
                document,
                numbered.group(1),
                style="List Number",
                num_id=active_number_id,
            )
            index += 1
            continue

        active_number_id = None

        marker = FIGURE_MARKER.fullmatch(stripped)
        if marker:
            _append_figure(document, root, marker.group(1), raster_cache, rasterizer)
            index += 1
            continue

        heading = HEADING.fullmatch(stripped)
        if heading:
            level = len(heading.group(1))
            document.add_paragraph(_plain(heading.group(2)), style=f"Heading {level}")
            index += 1
            continue

        if stripped.startswith("|"):
            headers, rows, index = _parse_table(lines, index)
            _append_table(document, headers, rows)
            continue

        if stripped.startswith("- "):
            _append_paragraph(document, stripped[2:], style="List Bullet")
            index += 1
            continue

        if stripped.startswith("图 "):
            paragraph = document.add_paragraph(_plain(stripped), style="Caption")
            paragraph.paragraph_format.keep_with_next = False
            index += 1
            continue

        if stripped.startswith("正式 artifact 为"):
            _append_paragraph(document, stripped, style="TxnMem Artifact")
        else:
            _append_paragraph(document, stripped)
        index += 1


def build_document(
    root: Path,
    output: Path,
    raster_cache: Path,
    rasterizer: Callable[[Path, Path], tuple[Path, float]] | None = None,
) -> Path:
    """Build, privacy-scrub and deterministically normalize the report package."""
    root = Path(root).resolve()
    output = Path(output)
    lines = (root / SOURCE).read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != f"# {TITLE}":
        raise ValueError("report source must begin with the approved title")
    metadata, body_start = _parse_cover(lines)
    selected_rasterizer = rasterizer or _rasterize_svg

    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _append_cover(document, metadata)
    _append_markdown(document, lines, body_start, root, Path(raster_cache), selected_rasterizer)

    properties = document.core_properties
    properties.title = TITLE
    properties.subject = "TxnMem paper-work and experiment evidence report"
    properties.author = ""
    properties.last_modified_by = ""
    properties.company = ""
    properties.comments = ""
    properties.keywords = "TxnMem; multi-agent memory; transactions; experiment evidence"

    output.parent.mkdir(parents=True, exist_ok=True)
    Path(raster_cache).mkdir(parents=True, exist_ok=True)
    document.save(output)
    _scrub_metadata(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--raster-cache", type=Path)
    args = parser.parse_args()
    if args.raster_cache is None:
        with tempfile.TemporaryDirectory(prefix="txnmem-report-raster-") as temporary:
            result = build_document(args.root, args.output, Path(temporary))
    else:
        result = build_document(args.root, args.output, args.raster_cache)
    print(result)


if __name__ == "__main__":
    main()
