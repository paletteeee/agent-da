#!/usr/bin/env python3
"""Deterministically build the anonymous TxnMem CCF-A manuscript DOCX.

Design preset: narrative_proposal_academic (named override of ``narrative_proposal``).
The constants below are intentionally data-like so Task 7 can audit the approved
page geometry, font fallbacks, academic header, table and accessibility decisions.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Task-6-auditable approved design data: narrative_proposal_academic.
DESIGN = {
    "preset": "narrative_proposal_academic",
    "page": {"size": "Letter", "margins_in": 1.0, "content_width_dxa": 9360,
             "header_footer_distance_in": 0.492},
    "fonts": {"body": "Calibri", "fallbacks": ["Arial Unicode MS", "Hiragino Sans GB"]},
    "body": {"size_pt": 11, "alignment": "justified", "after_pt": 8, "line_dxa": 320},
    "headings": {"h1": [16, "2E74B5", 18, 10], "h2": [13, "2E74B5", 12, 6],
                 "h3": [12, "1F4D78", 8, 4]},
    "lists": {"left_dxa": 540, "hanging_dxa": 280, "after_pt": 4, "line_dxa": 290},
    "tables": {"width_dxa": 9360, "indent_dxa": 120,
               "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
               "header_fill": "F4F6F9"},
    "captions": {"after_pt": 6, "size_pt": 10},
    "furniture": {"header": "TxnMem", "footer": "PAGE", "first_page_cover": False},
    "overrides": {
        "academic_title": "Chinese title + recommended English title + anonymous marker; no cover or border",
        "academic_header": "short title only; page-number footer; no commercial report chrome",
        "table_geometry": "all configured tables use fixed DXA widths, header repeat and padding",
    },
}

TITLE = "TxnMem：面向多 Agent 共享记忆的策略感知事务运行时"
ENGLISH_TITLE = "TxnMem: A Policy-Aware Transactional Runtime for Shared Memory in Multi-Agent Systems"
SHORT_TITLE = "TxnMem"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
MARKER = re.compile(r"^`?\[\[(FIG|TABLE):([a-z_]+)\]\]`?$")
HEADING = re.compile(r"^(#{1,3})\s+(.+?)\s*$")


def _load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _set_run_font(run, *, name: str = "Calibri", size: float | None = None,
                  color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    # Arial Unicode MS is retained as the approved fallback in DESIGN; Hiragino
    # is the available macOS CJK face and must be explicit for LibreOffice.
    rfonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_style_font(style, *, size: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Hiragino Sans GB")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def _spacing(paragraph_format, before: float, after: float, line_dxa: int) -> None:
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_dxa / 240


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(normal.paragraph_format, 0, 8, 320)

    for name, values in (("Heading 1", DESIGN["headings"]["h1"]),
                         ("Heading 2", DESIGN["headings"]["h2"]),
                         ("Heading 3", DESIGN["headings"]["h3"])):
        size, color, before, after = values
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        _spacing(style.paragraph_format, before, after, 240)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    _set_style_font(caption, size=10, color="404040")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(caption.paragraph_format, 4, 6, 240)
    caption.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("TxnMem Code", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(code, size=9.5, color="1F2937")
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    _spacing(code.paragraph_format, 3, 6, 220)

    references = doc.styles.add_style("TxnMem Reference", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(references, size=10.5)
    references.paragraph_format.left_indent = Inches(0.32)
    references.paragraph_format.first_line_indent = Inches(-0.32)
    _spacing(references.paragraph_format, 0, 4, 260)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in DESIGN["tables"]["cell_margins_dxa"].items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(DESIGN["tables"]["indent_dxa"]))
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _column_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    weights = [max(3, len(header), *(len(row[index]) for row in rows if index < len(row)))
               for index, header in enumerate(headers)]
    total = sum(weights)
    widths = [max(900, int(9360 * weight / total)) for weight in weights]
    widths[-1] += 9360 - sum(widths)
    return widths


def _add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.add_run(f"表 {len([p for p in doc.paragraphs if p.style.name == 'Caption' and p.text.startswith('表')]) + 1}  {title}")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        _set_cell_shading(cell, DESIGN["tables"]["header_fill"])
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, size=9.5, bold=True)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.add_run(value)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, size=9.25)
    _set_table_geometry(table, _column_widths(headers, rows))


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(header.add_run(SHORT_TITLE), size=9, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("第 "), size=9, color="666666")
    _add_page_field(footer)
    _set_run_font(footer.add_run(" 页"), size=9, color="666666")


def _plain(text: str) -> str:
    text = text.strip()
    text = text.replace("`", "")
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = text.replace("\\in", "∈")
    text = re.sub(r"\\[A-Za-z]+\{([^{}]+)\}", r"\1", text)
    text = text.replace("\\", "")
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def _parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        values = [_plain(value.strip()) for value in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows[0], rows[1:], index


def _add_list_item(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    _spacing(paragraph.paragraph_format, 0, 4, 290)
    paragraph.add_run(_plain(text))


def _rasterize_svg(svg: Path, destination: Path) -> Path:
    """Use macOS Quick Look deterministically because python-docx cannot embed SVG."""
    destination.mkdir(parents=True, exist_ok=True)
    # Quick Look may need an unsandboxed macOS XPC service. A digest-keyed shared
    # cache lets a preceding deterministic conversion be reused by test runners.
    digest = hashlib.sha256(svg.read_bytes()).hexdigest()[:16]
    shared = Path("/private/tmp/txnmem-docx-svg-cache") / f"{svg.stem}-{digest}.png"
    target = destination / f"{svg.stem}.png"
    if shared.is_file():
        shutil.copyfile(shared, target)
        return target
    with tempfile.TemporaryDirectory(prefix="txnmem-svg-", dir="/private/tmp") as tmp:
        tmp_path = Path(tmp)
        completed = subprocess.run(
            ["qlmanage", "-t", "-s", "2200", "-o", str(tmp_path), str(svg)],
            text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False,
        )
        generated = tmp_path / f"{svg.name}.png"
        if completed.returncode != 0 or not generated.is_file():
            raise RuntimeError(f"SVG rasterization failed for {svg.name}: {completed.stderr.strip()}")
        shared.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(generated, shared)
        shutil.copyfile(shared, target)
    return target


def _add_figure(doc: Document, figure_id: str, manifest: dict, root: Path, cache: Path) -> None:
    item = manifest["figures"][figure_id]
    png = _rasterize_svg(root / "paper_assets/figures" / item["file"], cache)
    ratio = item["dimensions"]["height"] / item["dimensions"]["width"]
    width = min(6.25, 6.0 / max(ratio, 0.5))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(png), width=Inches(width))
    shape._inline.docPr.set("descr", item["alt_text"])
    shape._inline.docPr.set("title", f"图：{figure_id}")
    caption = doc.add_paragraph(style="Caption")
    number = len([p for p in doc.paragraphs if p.style.name == "Caption" and p.text.startswith("图")]) + 1
    caption.add_run(f"图 {number}  {item['caption']}")


def _controlled_rows(root: Path) -> tuple[list[str], list[list[str]]]:
    data = _load(root / "results/paper_evidence/controlled_suite.json")["variants"]
    rows = []
    for name in ("TxnMem", "Naive", "TxnMem-NoTxn", "TxnMem-NoPolicyCommit", "TxnMem-NoRepair"):
        item = data[name]
        oracle = f"{item['oracle_match_count']}/400" if name != "TxnMem-NoRepair" else "未在此表报告"
        rows.append([name, f"{item['violation_count']}/400", oracle,
                     "受控套件的机制对照（详见正文）"])
    return ["变体", "违规 instance", "oracle 一致 instance", "机制解释"], rows


def _claim_ledger_rows(root: Path) -> tuple[list[str], list[list[str]]]:
    claims = _load(root / "configs/paper_claims.json")["claims"]
    active = [claim for claim in claims if claim.get("status") == "active"]
    return ["active claim", "evidence artifact", "claim boundary"], [
        [claim["claim_id"], claim["artifact_path"], claim["claim_boundary"]] for claim in active
    ]


TABLE_TITLES = {
    "requirements_gap": "设计需求与现有能力的缺口",
    "system_invariants": "TxnMem 的可检查不变量与主要机制",
    "workload_family": "受控 workload family 与历史转折",
    "experimental_setup": "评估层次、对象和判定器",
    "controlled_results": "受控套件中各实现变体的结果",
    "runtime_results": "模型、公开 runtime 与服务路径证据",
    "claim_ledger": "活跃 claim 与审计边界",
    "workload_schema": "评估记录的 schema 字段",
}


def _add_configured_table(doc: Document, table_id: str, markdown: dict[str, tuple[list[str], list[list[str]]]], root: Path) -> None:
    if table_id == "controlled_results":
        headers, rows = _controlled_rows(root)
    elif table_id == "claim_ledger":
        headers, rows = _claim_ledger_rows(root)
    else:
        headers, rows = markdown[table_id]
    _add_table(doc, TABLE_TITLES[table_id], headers, rows)


def _collect_markdown_tables(lines: list[str]) -> dict[str, tuple[list[str], list[list[str]]]]:
    found: dict[str, tuple[list[str], list[list[str]]]] = {}
    current: str | None = None
    index = 0
    while index < len(lines):
        marker = MARKER.match(lines[index].strip())
        if marker and marker.group(1) == "TABLE":
            current = marker.group(2)
        elif current and lines[index].lstrip().startswith("|"):
            headers, rows, index = _parse_table(lines, index)
            found[current] = (headers, rows)
            current = None
            continue
        index += 1
    return found


def _add_references(doc: Document, catalog: dict) -> None:
    for item in sorted(catalog["references"], key=lambda entry: entry["id"]):
        p = doc.add_paragraph(style="TxnMem Reference")
        p.add_run(
            f"[{item['id']}] " + "; ".join(item["authors"]) + ". "
            f"{item['title']}. {item['venue']}, {item['year']}. {item['url']}"
        )


def _append_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(title.paragraph_format, 0, 5, 240)
    _set_run_font(title.add_run(TITLE), size=17, color="0B2545", bold=True)
    english = doc.add_paragraph()
    english.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(english.paragraph_format, 0, 5, 240)
    _set_run_font(english.add_run(ENGLISH_TITLE), size=10.5, color="4B5563")
    anonymous = doc.add_paragraph()
    anonymous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(anonymous.paragraph_format, 0, 14, 240)
    _set_run_font(anonymous.add_run("匿名稿"), size=10.5, color="4B5563")


def _add_markdown(doc: Document, text: str, root: Path, manifest: dict, config: dict) -> None:
    lines = text.splitlines()
    markdown_tables = _collect_markdown_tables(lines)
    configured_tables = set(config["body_table_ids"] + config["appendix_table_ids"])
    cache = root / ".superpowers" / "sdd" / "2026-08-11-ccfa-paper-draft" / "docx_png_cache"
    index = 0
    in_references = False
    last_heading_level = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        heading = HEADING.match(stripped)
        if heading:
            title = heading.group(2)
            if title == TITLE:
                index += 1
                continue
            if title == "参考文献":
                doc.add_paragraph(title, style="Heading 1")
                last_heading_level = 1
                _add_references(doc, _load(root / "configs/txnmem_paper_references.json"))
                in_references = True
                index += 1
                continue
            if in_references and title != "附录":
                index += 1
                continue
            if title == "附录":
                in_references = False
            level = len(heading.group(1))
            # Markdown has a few related-work groups written directly as H3;
            # make the reader-facing Word outline accessible without changing prose.
            if level > last_heading_level + 1:
                level = last_heading_level + 1
            doc.add_paragraph(title, style=f"Heading {level}")
            last_heading_level = level
            index += 1
            continue
        if in_references or stripped == "匿名稿":
            index += 1
            continue
        marker = MARKER.match(stripped)
        if marker:
            kind, marker_id = marker.groups()
            if kind == "FIG":
                _add_figure(doc, marker_id, manifest, root, cache)
            elif marker_id in configured_tables:
                _add_configured_table(doc, marker_id, markdown_tables, root)
            index += 1
            while index < len(lines) and not lines[index].strip():
                index += 1
            if kind == "TABLE" and index < len(lines) and lines[index].lstrip().startswith("|"):
                _, _, index = _parse_table(lines, index)
            continue
        if stripped.startswith("图："):
            index += 1
            continue
        if stripped.startswith("```"):
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                doc.add_paragraph(_plain(lines[index]), style="TxnMem Code")
                index += 1
            index += 1
            continue
        if stripped.startswith("- "):
            _add_list_item(doc, stripped[2:])
            index += 1
            continue
        if stripped:
            doc.add_paragraph(_plain(stripped))
        index += 1


def _scrub_metadata(path: Path) -> None:
    """Strip creator/custom/revision identifiers while preserving generated OOXML."""
    with zipfile.ZipFile(path) as source:
        members = {name: source.read(name) for name in source.namelist() if name != "docProps/custom.xml"}
    core = ET.fromstring(members["docProps/core.xml"])
    for element in core:
        if element.tag.endswith("creator") or element.tag.endswith("lastModifiedBy"):
            element.text = ""
    members["docProps/core.xml"] = ET.tostring(core, encoding="utf-8", xml_declaration=True)
    for name, content in list(members.items()):
        if name.startswith("word/") and name.endswith(".xml"):
            root = ET.fromstring(content)
            for element in root.iter():
                for attribute in list(element.attrib):
                    if attribute.startswith(f"{{{W_NS}}}rsid"):
                        del element.attrib[attribute]
            members[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name in sorted(members):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            target.writestr(info, members[name])


def build_document(root: Path, output: Path) -> Path:
    """Build the final reader projection; author annotations are stripped first."""
    # This must remain the first operation on manuscript text: author claims never enter DOCX.
    from txnmem_manuscript_audit import strip_author_annotations

    root = root.resolve()
    config = _load(root / "configs/txnmem_ccfa_paper.json")
    raw_source = (root / config["paper_source_path"]).read_text(encoding="utf-8")
    reader_text = strip_author_annotations(raw_source)
    manifest = _load(root / "paper_assets/figures/manifest.json")

    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _append_title_block(doc)
    _add_markdown(doc, reader_text, root, manifest, config)
    properties = doc.core_properties
    properties.title = TITLE
    properties.subject = "Anonymous manuscript"
    properties.author = ""
    properties.last_modified_by = ""
    properties.company = ""
    properties.comments = ""
    properties.keywords = "TxnMem; multi-agent systems; shared memory; transactions"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    _scrub_metadata(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    print(build_document(args.root, args.output))


if __name__ == "__main__":
    main()
